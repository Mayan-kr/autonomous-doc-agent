"""Turn the agent's generated sections into a polished Word (.docx) file.

Uses python-docx. Handles headings, paragraphs, bullet lists, and simple
Markdown-style tables so the output looks like a real business document
rather than a wall of text.
"""

from __future__ import annotations

import os
import re
import uuid
from datetime import date
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")

# Splits text into runs while keeping the **bold**, *italic*, and `code` markers.
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
# Matches a Markdown heading line like "### Section".
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _add_inline(paragraph: Paragraph, text: str) -> None:
    """Add `text` to `paragraph`, converting inline Markdown to real runs.

    **bold** -> bold, *italic* -> italic, `code` -> monospace. Anything else is
    added as plain text, so stray markers never leak into the document.
    """
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def _add_title_block(doc: Document, title: str, doc_type: str, audience: str) -> None:
    """Cover heading: title, document type, audience, date."""
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"{doc_type}  •  Prepared for: {audience}  •  {date.today():%B %d, %Y}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer


def _render_body(doc: Document, text: str, section_title: str = "") -> None:
    """Render one section's body, interpreting light Markdown.

    Supported: '#'/'##'/'###' sub-headings, '- ' / '* ' bullets, '1.' numbered
    lists, '| a | b |' tables, and inline **bold** / *italic* / `code`.
    A heading line that just repeats `section_title` is dropped (the caller has
    already emitted the section heading), so titles never appear twice.
    """
    table_buffer: List[List[str]] = []

    def flush_table() -> None:
        if not table_buffer:
            return
        cols = max(len(r) for r in table_buffer)
        table = doc.add_table(rows=0, cols=cols)
        table.style = "Light Grid Accent 1"
        for row_cells in table_buffer:
            cells = table.add_row().cells
            for i in range(cols):
                _add_inline(cells[i].paragraphs[0], row_cells[i] if i < len(row_cells) else "")
        table_buffer.clear()
        doc.add_paragraph()

    for line in text.splitlines():
        stripped = line.strip()

        # Table row like: | Item | Cost |
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Skip the Markdown separator row (|---|---|)
            if set("".join(cells)) <= {"-", ":", " "}:
                continue
            table_buffer.append(cells)
            continue
        flush_table()

        if not stripped:
            continue

        # Markdown heading (### ...) -> a real Word sub-heading, nested under the
        # section's H1. Drop it if it merely repeats the section title.
        heading = _HEADING_RE.match(stripped)
        if heading:
            htext = heading.group(2).strip().strip("*").strip()
            if htext and htext.lower() != section_title.strip().lower():
                level = min(max(len(heading.group(1)), 2), 4)
                doc.add_heading(htext, level=level)
            continue

        if stripped.startswith(("- ", "* ")):
            _add_inline(doc.add_paragraph(style="List Bullet"), stripped[2:].strip())
        elif re.match(r"^\d+[.)]\s+", stripped):
            _add_inline(
                doc.add_paragraph(style="List Number"),
                re.sub(r"^\d+[.)]\s+", "", stripped),
            )
        else:
            _add_inline(doc.add_paragraph(), stripped)

    flush_table()


def build_document(
    *,
    title: str,
    doc_type: str,
    audience: str,
    assumptions: List[str],
    sections: List[Dict[str, str]],
) -> str:
    """Build the .docx and return its unique document id.

    `sections` is a list of {"title": ..., "content": ...} dicts.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()
    _add_title_block(doc, title, doc_type, audience)

    if assumptions:
        doc.add_heading("Assumptions", level=1)
        for item in assumptions:
            doc.add_paragraph(item, style="List Bullet")

    for section in sections:
        doc.add_heading(section["title"], level=1)
        _render_body(doc, section["content"], section_title=section["title"])

    document_id = uuid.uuid4().hex[:12]
    path = os.path.join(OUTPUT_DIR, f"{document_id}.docx")
    doc.save(path)
    return document_id


def path_for(document_id: str) -> str:
    """Absolute path for a generated document id."""
    return os.path.join(OUTPUT_DIR, f"{document_id}.docx")
